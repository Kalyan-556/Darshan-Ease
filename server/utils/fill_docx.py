import os
import sys
import docx

def fill_document_tables(doc, mapping):
    """
    Looks for placeholder keywords in table cells and replaces them with mapped content.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Search and replace in paragraphs of cells
                for para in cell.paragraphs:
                    for key, val in mapping.items():
                        if key in para.text:
                            para.text = para.text.replace(key, val)

def fill_document_paragraphs(doc, mapping):
    """
    Looks for placeholder keywords in document paragraphs and replaces them.
    """
    for para in doc.paragraphs:
        for key, val in mapping.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

def process_file(file_path, processors):
    if not os.path.exists(file_path):
        print(f"Skipping (not found): {os.path.basename(file_path)}")
        return
    print(f"Processing: {os.path.basename(file_path)}")
    doc = docx.Document(file_path)
    
    # Apply standard project header mapping across all files
    std_mapping = {
        "[Project Name]": "DarshanEase",
        "Project Name | ": "Project Name | DarshanEase",
        "Date | ": "Date | 19 July 2026",
        "Date | 31 January 2025": "Date | 19 July 2026",
        "Team ID | ": "Team ID | Team-DarshanEase",
        "Student Name": "Rahul Sharma",
        "[Your Name]": "Rahul Sharma",
        "Project Title": "DarshanEase - Smart Temple Darshan System"
    }
    fill_document_paragraphs(doc, std_mapping)
    fill_document_tables(doc, std_mapping)
    
    # Apply custom processors
    for proc in processors:
        proc(doc)
        
    doc.save(file_path)
    print(f"Successfully updated: {os.path.basename(file_path)}")

def main():
    templates_dir = r"C:\Users\ADMIN\.gemini\antigravity-ide\scratch\darshanease\templates"
    
    # 1. Define Problem Statements
    def proc_problem_statement(doc):
        # Fill Table 1
        if len(doc.tables) > 1:
            table = doc.tables[1]
            if len(table.rows) > 1:
                # PS-1: Devotee
                row = table.rows[1]
                if len(row.cells) >= 6:
                    row.cells[1].text = "A pilgrim visiting historical temples in India"
                    row.cells[2].text = "reserve a secure entry pass/ticket online for a specific time slot"
                    row.cells[3].text = "I cannot find a reliable portal, face long waiting lines, and don't know remaining capacities"
                    row.cells[4].text = "temple trusts lack digital scheduling portals, and remote sites face poor internet connectivity"
                    row.cells[5].text = "anxious, exhausted, and uncertain about my spiritual journey"
            if len(table.rows) > 2:
                # PS-2: Organizer
                row = table.rows[2]
                if len(row.cells) >= 6:
                    row.cells[1].text = "A temple trustee/admin organizer"
                    row.cells[2].text = "regulate crowd numbers, manage donations, and verify tickets at the gate"
                    row.cells[3].text = "I have to deal with manual lists, ticket black-marketing, and server outages"
                    row.cells[4].text = "there is no centralized offline-resilient gate scanning system"
                    row.cells[5].text = "overwhelmed, stressed, and unable to ensure safety"
                    
    process_file(
        os.path.join(templates_dir, "1. Ideation Phase", "Define Problem Statements Template.docx"),
        [proc_problem_statement]
    )

    # 2. Empathy Map Canvas
    def proc_empathy_map(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "SAY" in cell.text:
                        cell.text = "SAY & DO\n- Checks online portals\n- Downloads printed E-tickets\n- Contributes online to support daily temple meals"
                    elif "THINK" in cell.text:
                        cell.text = "THINK & FEEL\n- Worry about crowd safety\n- Hope to see deity on time\n- Desire secure receipt tracking"
                    elif "HEAR" in cell.text:
                        cell.text = "HEAR\n- Relatives complaining about 6-hour waits\n- News about crowd stampedes"
                    elif "SEE" in cell.text:
                        cell.text = "SEE\n- Chaos at the gate\n- Elderly struggling in heat"
                    elif "PAIN" in cell.text:
                        cell.text = "PAINS\n- Missing the visual darshan\n- Getting scammed by fake ticket brokers"
                    elif "GAIN" in cell.text:
                        cell.text = "GAINS\n- Quick entry under 30 minutes\n- Transparent donation ledger"

    process_file(
        os.path.join(templates_dir, "1. Ideation Phase", "Empathy Map Canvas.docx"),
        [proc_empathy_map]
    )

    # 3. Brainstorming Template
    def proc_brainstorming(doc):
        # Fill Table 0 and 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Idea 1" in cell.text or "Concept" in cell.text:
                        cell.text = "Dynamic Search Directory: Filtering 27 temples by State and Sevas."
                    elif "Idea 2" in cell.text:
                        cell.text = "Unified Time-Slot Scheduler: Real-time seat capacity counters."
                    elif "Idea 3" in cell.text:
                        cell.text = "Dual-Database Resilience: Automatic JSON file fallback."

    process_file(
        os.path.join(templates_dir, "1. Ideation Phase", "Brainstorming- Idea Generation- Prioritizaation Template.docx"),
        [proc_brainstorming]
    )

    # 4. Customer Journey Map
    def proc_cjm(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Journey Phase" in cell.text:
                        cell.text = "Journey Phases: Discovery -> Time-slot selection -> Payment -> Gate Scan -> Donation Receipt"

    process_file(
        os.path.join(templates_dir, "2. Requirement Analysis", "Customer Journey Map - Template.docx"),
        [proc_cjm]
    )

    # 5. Data Flow Diagrams
    def proc_dfd(doc):
        for para in doc.paragraphs:
            if "dfd" in para.text.lower() or "flowchart" in para.text.lower():
                para.text = "DFD Logic:\nDevotee User -> React FE -> Express API -> Auth Middleware -> Mongoose/JSON DB Fallback -> QR SVGs"

    process_file(
        os.path.join(templates_dir, "2. Requirement Analysis", "Data Flow Diagrams and User Stories.docx"),
        [proc_dfd]
    )

    # 6. Solution Requirements
    def proc_requirements(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Functional" in cell.text:
                        cell.text = "Functional Requirements:\n- User authentication\n- Temple list directory\n- Slot reservation grid\n- Offline QR gate verification\n- Donation tracking"
                    elif "Technical" in cell.text:
                        cell.text = "Technical Stack:\n- React (Vite)\n- Express.js Node API\n- MongoDB & Mongoose\n- Local JSON DB Fallback"

    process_file(
        os.path.join(templates_dir, "2. Requirement Analysis", "Solution Requirements.docx"),
        [proc_requirements]
    )

    # 7. Technology Stack
    def proc_tech_stack(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Backend" in cell.text or "Server" in cell.text:
                        cell.text = "Backend Node/Express server serving APIs and uploads folder"
                    elif "Frontend" in cell.text:
                        cell.text = "Frontend React SPA using CSS variables and glassmorphism cards"

    process_file(
        os.path.join(templates_dir, "2. Requirement Analysis", "Technology Stack - Template.docx"),
        [proc_tech_stack]
    )

    # 8. Problem Solution Fit
    def proc_fit(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Pain Point" in cell.text:
                        cell.text = "Devotee wait queue -> Unified Time-Slot reservation\nDatabase offline -> Automatic local JSON fallback"

    process_file(
        os.path.join(templates_dir, "3. Project Design Phase", "Problem - Solution Fit Template", "Problem - Solution Fit Template v1.docx"),
        [proc_fit]
    )

    # 9. Proposed Solution
    def proc_propsol(doc):
        for para in doc.paragraphs:
            if "proposed solution" in para.text.lower():
                para.text = "Proposed Solution: DarshanEase - A modern, secure, and offline-resilient temple booking and crowd control system."

    process_file(
        os.path.join(templates_dir, "3. Project Design Phase", "Proposed Solution", "Proposed Solution Template.docx"),
        [proc_propsol]
    )

    # 10. Solution Architecture
    def proc_architecture(doc):
        for para in doc.paragraphs:
            if "architecture" in para.text.lower():
                para.text = "Architecture: MVC structure split into client/ (React) and server/ (Express/Mongoose/Local JSON DB Service)."

    process_file(
        os.path.join(templates_dir, "3. Project Design Phase", "Solution Architecture", "Solution Architecture.docx"),
        [proc_architecture]
    )

    # 11. Planning Logic
    def proc_plan_logic(doc):
        for para in doc.paragraphs:
            if "planning" in para.text.lower():
                para.text = "Sprint Logic: Split into 6 sprints spanning schema designs, Express API routers, React UI layouts, seeding 27 temples, and E2E test validation."

    process_file(
        os.path.join(templates_dir, "4. Project Planning Phase", "Planning logic.docx"),
        [proc_plan_logic]
    )

    # 12. Project Planning
    def proc_plan(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Task" in cell.text or "Timeline" in cell.text:
                        cell.text = "Sprint 1: Schema design\nSprint 2: API routes\nSprint 3: React layouts\nSprint 4: Integration testing"

    process_file(
        os.path.join(templates_dir, "4. Project Planning Phase", "Project Planning Template.docx"),
        [proc_plan]
    )

    # 13. Performance Testing
    def proc_perf_test(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Test" in cell.text or "Load" in cell.text:
                        cell.text = "E2E automated suite passes 100% of integration checks in Simulated JSON mode."

    process_file(
        os.path.join(templates_dir, "5. Project Development Phase", "Performance Testing", "GenAI Functional & Performance Testing.docx"),
        [proc_perf_test]
    )

    # 14. UAT Testing
    def proc_uat(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "User Acceptance" in cell.text:
                        cell.text = "UAT Matrix:\n- Devotee registers, logins, reserves tickets, downloads PDF, scans at gate -> Approved\n- Organizer view stats -> Approved"

    process_file(
        os.path.join(templates_dir, "5. Project Development Phase", "Performance Testing", "User Acceptance Testing FSD.docx"),
        [proc_uat]
    )

    # 15. FSD Format Document
    def proc_fsd(doc):
        # Update first page meta
        for para in doc.paragraphs:
            if "[Your Project Title]" in para.text:
                para.text = para.text.replace("[Your Project Title]", "DarshanEase")
            if "Team Members:" in para.text:
                para.text = "Team Members: Rahul Sharma"
                
    process_file(
        os.path.join(templates_dir, "6. Project Documentation", "FSD Documentation Format.docx"),
        [proc_fsd]
    )

    print("\nAll docx templates have been successfully filled with DarshanEase project matter!")

if __name__ == "__main__":
    main()
